"""Post-generation output validation."""
from __future__ import annotations
import re
from pathlib import Path
from docx import Document


def validate_output(docx_path: str, expected_profiles: int, expected_rates: list[int]) -> list[str]:
    """Validate the generated DOCX. Returns list of errors."""
    errors = []

    try:
        doc = Document(docx_path)
    except Exception as e:
        return [f"Cannot open generated DOCX: {e}"]

    errors.extend(_check_no_tokens(doc))
    errors.extend(_check_table1(doc, expected_profiles, expected_rates))
    errors.extend(_check_section7(doc, expected_profiles))
    errors.extend(_check_revision_record(doc))
    errors.extend(_check_protection(docx_path))

    return errors


def _check_no_tokens(doc: Document) -> list[str]:
    errors = []
    for i, p in enumerate(doc.paragraphs):
        txt = "".join(r.text for r in p.runs)
        if "{{" in txt and "}}" in txt:
            errors.append(f"Unresolved token in paragraph {i}: {txt[:80]}")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = "".join(r.text for r in p.runs)
                    if "{{" in txt and "}}" in txt:
                        errors.append(f"Unresolved token in table cell: {txt[:80]}")
    return errors


def _check_table1(doc: Document, expected_profiles: int, expected_rates: list[int]) -> list[str]:
    errors = []
    table = _find_table1(doc)
    if table is None:
        errors.append("Duty Profile Test Matrix table not found")
        return errors

    data_rows = len(table.rows) - 1
    expected_total = sum(expected_rates)
    if data_rows != expected_total:
        errors.append(f"Table 1 has {data_rows} data rows, expected {expected_total}")

    return errors


def _check_section7(doc: Document, expected_blocks: int) -> list[str]:
    errors = []
    block_count = 0
    for p in doc.paragraphs:
        txt = p.text
        if re.match(r"7\.\d+\s+Qualification Tests", txt):
            block_count += 1

    if block_count != expected_blocks:
        errors.append(f"Section 7 has {block_count} blocks, expected {expected_blocks}")

    return errors


def _check_revision_record(doc: Document) -> list[str]:
    errors = []
    for table in doc.tables:
        if len(table.columns) == 4:
            header = "".join(c.text for c in table.rows[0].cells)
            if "Revision" in header and "Date" in header and "Description" in header:
                if len(table.rows) > 1:
                    row1 = table.rows[1]
                    if row1.cells[0].text.strip() != "00":
                        errors.append(f"Revision Record first cell is '{row1.cells[0].text}', expected '00'")
                    for c in row1.cells[1:]:
                        if c.text.strip():
                            errors.append(f"Revision Record should be blank, but cell has: {c.text}")
    return errors


def _check_protection(docx_path: str) -> list[str]:
    errors = []
    import zipfile
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            if "word/settings.xml" in z.namelist():
                settings = z.read("word/settings.xml").decode("utf-8")
                if "documentProtection" not in settings:
                    errors.append("No document protection found in settings.xml")
            else:
                errors.append("settings.xml not found in DOCX")

            if "word/document.xml" in z.namelist():
                doc_xml = z.read("word/document.xml").decode("utf-8")
                perm_starts = doc_xml.count("permStart")
                perm_ends = doc_xml.count("permEnd")
                if perm_starts == 0:
                    errors.append("No permStart markers found in document.xml")
                if perm_starts != perm_ends:
                    errors.append(f"permStart ({perm_starts}) != permEnd ({perm_ends})")
    except Exception as e:
        errors.append(f"Cannot check protection: {e}")
    return errors


def _find_table1(doc: Document):
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.columns) == 6:
            header = "".join(c.text for c in table.rows[0].cells)
            if "Duty Profile" in header and "Conditioning" in header:
                return table
    return None
