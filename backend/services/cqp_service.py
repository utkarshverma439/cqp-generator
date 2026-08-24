"""CQP Service - orchestrates the entire generation pipeline."""
from __future__ import annotations
import shutil
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from backend.parsers.tmp_parser import parse_tmp
from backend.parsers.acl_parser import parse_acl
from backend.parsers.datasheet_parser import parse_datasheet
from backend.services.normalizer import normalize
from backend.generator.token_replacer import build_token_map, replace_simple_tokens
from backend.generator.table1_generator import generate_table1
from backend.generator.section7_generator import generate_section7
from backend.generator.footnote_handler import add_footnotes
from backend.generator.protection import apply_document_protection
from backend.validators.input_validator import validate_input
from backend.validators.output_validator import validate_output
from backend.config import TEMPLATE_PATH


def generate_cqp(tmp_path: str, acl_path: str, ds_path: str, market: str,
                 output_dir: str = "outputs") -> dict:
    """Generate a CQP document from source files.

    Returns dict with 'output_path' on success or 'errors' on failure.
    """
    tmp_data = parse_tmp(tmp_path)
    acl_data = parse_acl(acl_path)
    ds_data = parse_datasheet(ds_path)

    cell_data = normalize(tmp_data, acl_data, ds_data, market)

    input_errors = validate_input(cell_data)
    if input_errors:
        return {"success": False, "error": "Input validation failed", "details": input_errors}

    output_path = _generate_docx(cell_data, output_dir)

    expected_rates = [len(p.conditioning_rates) for p in cell_data.duty_profiles]
    output_errors = validate_output(
        str(output_path),
        expected_profiles=len(cell_data.duty_profiles),
        expected_rates=expected_rates,
    )
    if output_errors:
        return {"success": False, "error": "Output validation failed", "details": output_errors}

    return {"success": True, "output_path": str(output_path), "cell_data": cell_data}


def _generate_docx(cell_data, output_dir: str) -> Path:
    """Generate the DOCX from CellData."""
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    output_path = out_dir / f"Cell_Qualification_Protocol_{cell_data.cell_model}.docx"

    shutil.copy2(str(TEMPLATE_PATH), str(output_path))
    doc = Document(str(output_path))

    token_map = build_token_map(cell_data)
    replace_simple_tokens(doc, token_map)

    _remove_template_instructions(doc)
    _remove_reviewer_token_in_section7(doc)

    generate_table1(doc, cell_data)
    generate_section7(doc, cell_data)
    add_footnotes(doc, cell_data)

    doc.save(str(output_path))
    apply_document_protection(str(output_path))

    return output_path


def _remove_template_instructions(doc: Document) -> None:
    """Remove template instruction paragraphs that are not real content."""
    instructions = [
        "The following values are taken from the vendor datasheet.",
        "Table 1 \u2014 expands to one row per duty profile and conditioning rate.",
        "The block below is repeated for each duty profile.",
        "This table is completed manually and MUST remain blank at issue.",
    ]
    body = doc.element.body
    for p in list(body.iter(qn("w:p"))):
        txt = "".join(t.text for t in p.iter() if t.tag.endswith("}t") and t.text)
        for inst in instructions:
            if inst in txt:
                parent = p.getparent()
                if parent is not None:
                    parent.remove(p)
                break


def _remove_reviewer_token_in_section7(doc: Document) -> None:
    """Replace {{ to_be_added_by_reviewer }} in Section 7 with empty string
    for the labels, but the actual content paragraphs use boilerplate.
    """
    from backend.config import ACCEPTANCE_CRITERIA_TEXT, CONCLUSION_TEXT

    found_acceptance = False
    found_conclusion = False

    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt == "Acceptance Criteria:" or txt == "{{ to_be_added_by_reviewer }}":
            if not found_acceptance and "{{ to_be_added_by_reviewer }}" in txt:
                for run in p.runs:
                    run.text = run.text.replace("{{ to_be_added_by_reviewer }}", ACCEPTANCE_CRITERIA_TEXT)
                found_acceptance = True
            elif "{{ to_be_added_by_reviewer }}" in txt:
                for run in p.runs:
                    run.text = run.text.replace("{{ to_be_added_by_reviewer }}", "")
        elif txt == "Conclusion:":
            found_conclusion = False
        elif "{{ to_be_added_by_reviewer }}" in txt and not found_acceptance:
            for run in p.runs:
                run.text = run.text.replace("{{ to_be_added_by_reviewer }}", ACCEPTANCE_CRITERIA_TEXT)
            found_acceptance = True
        elif "{{ to_be_added_by_reviewer }}" in txt and found_acceptance and not found_conclusion:
            for run in p.runs:
                run.text = run.text.replace("{{ to_be_added_by_reviewer }}", CONCLUSION_TEXT)
            found_conclusion = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if "{{ to_be_added_by_reviewer }}" in run.text:
                            run.text = run.text.replace("{{ to_be_added_by_reviewer }}", "")
