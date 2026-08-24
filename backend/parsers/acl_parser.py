"""ACL (Acceptance Criteria & Limits) parser.

Extracts duty profiles, conditioning rates, tests, limits, cycle counts,
markers, and footnotes from ACL documents. The parser is dynamic and does
not assume a fixed number of profiles or tests.
"""
from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.table import Table

from backend.models.cell_data import ACLData, DutyProfile, TestEntry, Footnote


# Regex to detect duty profile headings
# Matches: "Duty Profile: Automotive Traction   (conditioning rates: 0.5C, 1.0C)"
# Also:    "Duty Profile: Aerospace Auxiliary   (conditioning rates: 0.33C, 0.5C)"
PROFILE_RE = re.compile(
    r'Duty\s+Profile:\s*(.+?)\s*\(conditioning\s+rates?:\s*(.+?)\)',
    re.IGNORECASE
)

# Regex to extract cycle count from Cycle-Life Endurance acceptance limit
# Matches: "after 500 cycles", "after 800 cycles"
CYCLE_COUNT_RE = re.compile(r'after\s+(\d+)\s+cycles', re.IGNORECASE)


def parse_acl(file_path: str | Path) -> ACLData:
    """Parse an ACL .docx file and extract all structured data."""
    doc = Document(str(file_path))

    title = _extract_title(doc)
    doc_id, manufacturer = _extract_metadata(doc)
    profiles = _extract_profiles(doc)
    footnotes = _extract_footnotes(doc)

    if not profiles:
        raise ValueError("No duty profiles found in ACL document")

    return ACLData(
        title=title,
        doc_id=doc_id,
        manufacturer=manufacturer,
        duty_profiles=profiles,
        footnotes=footnotes,
    )


def _extract_title(doc: Document) -> str:
    """Extract the ACL document title (first paragraph)."""
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            return text
    raise ValueError("Could not find ACL title")


def _extract_metadata(doc: Document) -> tuple[str, str]:
    """Extract document ID and manufacturer from the metadata line.

    The metadata line is typically the second paragraph, italic, containing:
    '<manufacturer>  |  Document ACL-XXXXXX-00  |  Framework IESF-4400'
    """
    for p in doc.paragraphs:
        text = p.text.strip()
        if '|' in text:
            parts = [part.strip() for part in text.split('|')]
            manufacturer = parts[0] if parts else ""
            doc_id = ""
            for part in parts:
                part_stripped = part.strip()
                if part_stripped.upper().startswith('ACL-') or part_stripped.upper().startswith('DOCUMENT'):
                    # Remove "Document " prefix if present
                    doc_id = re.sub(r'^Document\s+', '', part_stripped, flags=re.IGNORECASE)
                    break
            return doc_id, manufacturer

    return "", ""


def _extract_profiles(doc: Document) -> list[DutyProfile]:
    """Extract all duty profiles with their tests from the ACL.

    Strategy:
    1. Walk through all body elements (paragraphs and tables)
    2. Detect profile headings via regex
    3. Associate each profile heading with the next table
    4. Parse the table for test data
    """
    profiles = []

    # Build a list of body elements in order
    body = doc.element.body
    elements = list(body)

    current_profile_name = None
    current_rates = None

    for elem in elements:
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag

        if tag == 'p':
            # It's a paragraph - check if it's a profile heading
            text = elem.text or ''
            # Also check runs for full text
            for r in elem.iter():
                if r.tag.endswith('}t') and r.text:
                    pass  # text already captured via elem.text
            # Get full text from all w:t elements
            full_text = ''.join(
                t.text for t in elem.iter() if t.tag.endswith('}t') and t.text
            )

            match = PROFILE_RE.search(full_text)
            if match:
                # Save previous profile if any
                if current_profile_name is not None:
                    pass  # will be saved when we find the table

                current_profile_name = match.group(1).strip()
                rates_str = match.group(2).strip()
                current_rates = [r.strip() for r in rates_str.split(',')]
                continue

        elif tag == 'tbl':
            # It's a table - if we have a pending profile, this is its test table
            if current_profile_name is not None:
                table = Table(elem, doc)
                tests = _parse_test_table(table)
                cycle_count = _extract_cycle_count(tests)

                profile = DutyProfile(
                    name=current_profile_name,
                    conditioning_rates=current_rates or [],
                    cycle_count=cycle_count,
                    tests=tests,
                )
                profiles.append(profile)
                current_profile_name = None
                current_rates = None

    return profiles


def _parse_test_table(table: Table) -> list[TestEntry]:
    """Parse a test table into TestEntry objects.

    Table structure:
    Header: Sr. No. | Test Parameter | Acceptance Limit | IESF-4400 Clause
    Data rows: 1 | Capacity Verification | >= 98.0% ... * $ | §6.2
    """
    tests = []

    # Skip header row (row 0)
    for i, row in enumerate(table.rows[1:], start=1):
        cells = [cell.text.strip() for cell in row.cells]

        if len(cells) < 4:
            continue

        sr_no_str = cells[0]
        test_name = cells[1]
        acceptance_limit = cells[2]
        clause = cells[3]

        # Skip empty rows
        if not test_name:
            continue

        try:
            sr_no = int(sr_no_str)
        except (ValueError, TypeError):
            sr_no = i

        tests.append(TestEntry(
            sr_no=sr_no,
            test_name=test_name,
            acceptance_limit=acceptance_limit,
            clause=clause,
        ))

    return tests


def _extract_cycle_count(tests: list[TestEntry]) -> int:
    """Extract cycle count from the Cycle-Life Endurance test row."""
    for test in tests:
        if 'cycle' in test.test_name.lower() and 'life' in test.test_name.lower():
            match = CYCLE_COUNT_RE.search(test.acceptance_limit)
            if match:
                return int(match.group(1))

    raise ValueError(
        "Could not extract cycle count from Cycle-Life Endurance test. "
        f"Tests found: {[t.test_name for t in tests]}"
    )


def _extract_footnotes(doc: Document) -> list[Footnote]:
    """Extract footnotes from the ACL document.

    Footnotes appear after all profile tables, typically starting with
    a 'Notes:' heading followed by lines like:
        * Capacity limits apply to ...
        # DCIR is measured by ...
    """
    footnotes = []

    # Find the "Notes:" heading
    in_notes = False
    for p in doc.paragraphs:
        text = p.text.strip()

        if text.lower() in ('notes:', 'note:', 'notes'):
            in_notes = True
            continue

        if in_notes and text:
            # Check if this line starts with a marker character
            marker = text[0]
            if marker in ('*', '#', '@', '$'):
                note_text = text[1:].strip()
                footnotes.append(Footnote(marker=marker, text=note_text))

    return footnotes
