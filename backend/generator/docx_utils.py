"""Low-level OOXML utilities for DOCX manipulation."""
from __future__ import annotations
import copy
from lxml import etree
from docx import Document
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.table import Table, _Row


def replace_tokens_in_paragraph(paragraph, token_map: dict[str, str]) -> None:
    """Replace {{ token }} patterns in a paragraph while preserving formatting."""
    full_text = "".join(run.text for run in paragraph.runs)
    if "{{" not in full_text:
        return

    new_text = full_text
    for token, value in token_map.items():
        new_text = new_text.replace("{{ " + token + " }}", value)
        new_text = new_text.replace("{{" + token + "}}", value)

    if new_text != full_text and paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def replace_tokens_in_table(table, token_map: dict[str, str]) -> None:
    """Replace tokens in all cells of a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_tokens_in_paragraph(paragraph, token_map)


def replace_tokens_in_header_footer(doc: Document, token_map: dict[str, str]) -> None:
    """Replace tokens in headers and footers."""
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_tokens_in_paragraph(paragraph, token_map)
        for paragraph in section.footer.paragraphs:
            replace_tokens_in_paragraph(paragraph, token_map)


def clone_table_row(table: Table, source_row_index: int) -> _Row:
    """Clone a table row by deep-copying its XML element."""
    source_row = table.rows[source_row_index]
    new_tr = copy.deepcopy(source_row._tr)
    table._tbl.append(new_tr)
    return _Row(new_tr, table)


def set_cell_vmerge(cell, val: str = "restart") -> None:
    """Set vertical merge on a table cell via OOXML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:vMerge")):
        tcPr.remove(existing)
    vMerge = OxmlElement("w:vMerge")
    vMerge.set(qn("w:val"), val)
    tcPr.append(vMerge)


def remove_cell_vmerge(cell) -> None:
    """Remove vertical merge from a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:vMerge")):
        tcPr.remove(existing)


def set_cell_shading(cell, fill_color: str) -> None:
    """Set cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, font_size: int = 20,
                  font_color: str = "000000", alignment: str = "left") -> None:
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = font_size * 12700  # half-points to EMU
    if font_color != "000000":
        rPr = run._r.get_or_add_rPr()
        color = OxmlElement("w:color")
        color.set(qn("w:val"), font_color)
        rPr.append(color)
    if alignment:
        pPr = p._p.get_or_add_pPr()
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), alignment)
        pPr.append(jc)


def copy_row_formatting(source_row, target_row) -> None:
    """Copy cell formatting from source row to target row."""
    for src_cell, tgt_cell in zip(source_row.cells, target_row.cells):
        src_tc = src_cell._tc
        tgt_tc = tgt_cell._tc
        src_tcPr = src_tc.find(qn("w:tcPr"))
        if src_tcPr is not None:
            new_tcPr = copy.deepcopy(src_tcPr)
            tgt_tcPr = tgt_tc.find(qn("w:tcPr"))
            if tgt_tcPr is not None:
                tgt_tc.remove(tgt_tcPr)
            tgt_tc.insert(0, new_tcPr)


def add_permission_marker(element, perm_id: int, edit: str = "everyone") -> None:
    """Add w:permStart before an element."""
    perm_start = OxmlElement("w:permStart")
    perm_start.set(qn("w:id"), str(perm_id))
    perm_start.set(qn("w:edit"), edit)
    element.addprevious(perm_start)


def add_permission_end(element, perm_id: int) -> None:
    """Add w:permEnd after an element."""
    perm_end = OxmlElement("w:permEnd")
    perm_end.set(qn("w:id"), str(perm_id))
    element.addnext(perm_end)


def find_paragraph_with_text(doc: Document, text: str):
    """Find a paragraph containing the given text."""
    for p in doc.paragraphs:
        if text in p.text:
            return p
    return None
