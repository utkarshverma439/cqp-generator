"""Table 1 (Duty Profile Test Matrix) generator.

Expands the template's placeholder row into one row per
duty_profile x conditioning_rate, with proper vMerge.
"""
from __future__ import annotations
import copy
from docx import Document
from docx.table import Table
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from backend.models.cell_data import CellData
from backend.generator.docx_utils import set_cell_vmerge


def generate_table1(doc: Document, data: CellData) -> None:
    """Expand the Duty Profile Test Matrix table."""
    table = _find_table1(doc)
    if table is None:
        raise ValueError("Could not find Duty Profile Test Matrix table")

    placeholder_row_idx = 1
    placeholder_row = table.rows[placeholder_row_idx]

    total_rows = sum(len(p.conditioning_rates) for p in data.duty_profiles)
    if total_rows == 0:
        return

    # Clone the placeholder row's XML to use as a template
    template_tr = copy.deepcopy(placeholder_row._tr)

    # Remove the original placeholder row
    table._tbl.remove(placeholder_row._tr)

    # Create all data rows from template
    all_data_rows = []
    for _ in range(total_rows):
        new_tr = copy.deepcopy(template_tr)
        table._tbl.append(new_tr)
        from docx.table import _Row
        all_data_rows.append(_Row(new_tr, table))

    sr_no = 1
    row_idx = 0
    for profile in data.duty_profiles:
        for rate_idx, rate in enumerate(profile.conditioning_rates):
            row = all_data_rows[row_idx]
            cells = row.cells

            cells[0].text = str(sr_no)
            cells[1].text = profile.name if rate_idx == 0 else ""
            cells[2].text = rate
            cells[3].text = data.electrical_ratings.v_max
            cells[4].text = data.electrical_ratings.v_min
            cells[5].text = str(profile.cycle_count)

            for cell in cells:
                for p in cell.paragraphs:
                    p.alignment = None
                    for run in p.runs:
                        run.bold = False
                        run.font.size = 10 * 12700
                        rPr = run._r.get_or_add_rPr()
                        for elem in list(rPr.findall(qn("w:b"))):
                            rPr.remove(elem)

            tc = cells[1]._tc
            tcPr = tc.get_or_add_tcPr()
            for existing in tcPr.findall(qn("w:vMerge")):
                tcPr.remove(existing)

            if len(profile.conditioning_rates) > 1:
                if rate_idx == 0:
                    set_cell_vmerge(cells[1], "restart")
                else:
                    set_cell_vmerge(cells[1], "continue")
            else:
                vMerge = OxmlElement("w:vMerge")
                vMerge.set(qn("w:val"), "restart")
                tcPr.append(vMerge)

            sr_no += 1
            row_idx += 1

    _apply_data_row_formatting(table)
    _add_table1_caption(doc, table)


def _find_table1(doc: Document) -> Table | None:
    """Find the Duty Profile Test Matrix table (the one with 'Sr. No.' and 'Duty Profile' headers)."""
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.columns) == 6:
            header_text = "".join(cell.text for cell in table.rows[0].cells)
            if "Duty Profile" in header_text and "Conditioning" in header_text:
                return table
    return None


def _apply_data_row_formatting(table: Table) -> None:
    """Apply consistent formatting to all data rows."""
    for i, row in enumerate(table.rows):
        if i == 0:
            continue
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is not None:
                shd = tcPr.find(qn("w:shd"))
                if shd is not None:
                    tcPr.remove(shd)

            for p in cell.paragraphs:
                pPr = p._p.get_or_add_pPr()
                for jc in list(pPr.findall(qn("w:jc"))):
                    pPr.remove(jc)
                for run in p.runs:
                    rPr = run._r.get_or_add_rPr()
                    for b in list(rPr.findall(qn("w:b"))):
                        if b.get(qn("w:val")) == "false":
                            pass
                        else:
                            rPr.remove(b)
                    b_elem = OxmlElement("w:b")
                    b_elem.set(qn("w:val"), "false")
                    rPr.append(b_elem)
                    run.font.size = 10 * 12700


def _add_table1_caption(doc: Document, table: Table) -> None:
    """Add 'Table 1. Duty profile test matrix.' paragraph after the table."""
    tbl_elem = table._tbl
    tbl_parent = tbl_elem.getparent()
    tbl_index = list(tbl_parent).index(tbl_elem)

    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    b = OxmlElement("w:b")
    rPr.append(b)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "Table 1. Duty profile test matrix."
    r.append(t)
    p.append(r)
    tbl_parent.insert(tbl_index + 1, p)
